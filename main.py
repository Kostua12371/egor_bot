from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, ContextTypes, filters
from docx import Document
import docx.shared
import os
from datetime import datetime
from dotenv import load_dotenv
import os

load_dotenv()
TOKEN = os.getenv("TELEGRAM_TOKEN")

field_labels = [
    "Размеры:",
    "Количество макетов:",
    "Количество наклеек:",
    "Тип резки:",
    "Тип наклеек:",
    "Имя заказчика:",
    "Сумма заказа:",
    "Себестоимость:",
    "Адрес доставки:"
]

user_data = {}
order_count = 0


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_data[update.effective_chat.id] = {
        'step': -1,
        'data': {}
    }
    await update.message.reply_text("🛒 Новый заказ\n\n📸 Пришлите изображение для заказа")


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_chat.id
    if user_id not in user_data:
        await update.message.reply_text("Начните с команды /start")
        return

    # Сохраняем фотографию
    photo = await update.message.photo[-1].get_file()
    photo_path = f'temp_{user_id}.jpg'
    await photo.download_to_drive(photo_path)
    user_data[user_id]['data']['photo_path'] = photo_path
    user_data[user_id]['step'] = 0

    await update.message.reply_text(f"1️⃣ {field_labels[0]}")


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_chat.id
    if user_id not in user_data:
        await update.message.reply_text("Начните с команды /start")
        return

    current_step = user_data[user_id]['step']

    if current_step < len(field_labels):

        user_data[user_id]['data'][field_labels[current_step]] = update.message.text
        user_data[user_id]['step'] += 1

        if current_step + 1 < len(field_labels):
            await update.message.reply_text(f"➡️ {field_labels[current_step + 1]}")
        else:

            await create_document(update, user_data[user_id]['data'])
            user_data.pop(user_id)


async def create_document(update: Update, data):
    global order_count
    order_count += 1

    doc = Document()
    doc.add_heading(f'Заказ #{order_count}', level=1)

    if 'photo_path' in data:
        try:
            doc.add_picture(data['photo_path'], width=docx.shared.Inches(4))
            doc.add_paragraph("\n")
        except Exception as e:
            print(f"Ошибка при добавлении фото: {e}")

    for label, value in data.items():
        if label != 'photo_path':
            doc.add_paragraph(f"🔹 {label} {value}")

    doc.add_paragraph(f"\n📅 Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    doc_path = f'order_{update.effective_chat.id}.docx'
    doc.save(doc_path)

    with open(doc_path, 'rb') as file:
        await update.message.reply_document(document=file, caption="✅ Ваш заказ готов!")

    if 'photo_path' in data:
        os.remove(data['photo_path'])
    os.remove(doc_path)

    await update.message.reply_text("Чтобы сделать новый заказ, отправь /start")


if __name__ == '__main__':
    app = Application.builder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))

    app.run_polling()
